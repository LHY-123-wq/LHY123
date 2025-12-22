import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn

# 读取Markdown文件
with open('数字化转型指数查询App开发报告.md', 'r', encoding='utf-8') as f:
    markdown_content = f.read()

# 创建Word文档
doc = Document()

# 设置文档样式
doc.styles['Normal'].font.name = 'Microsoft YaHei'
doc.styles['Normal'].font.size = Pt(12)

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(level=level)
    heading_run = heading.add_run(text)
    heading_run.font.name = 'Microsoft YaHei'
    heading_run.font.size = Pt(16 if level == 0 else 14)
    return heading

def add_paragraph(doc, text):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(12)
    return para

def add_code_block(doc, code):
    """添加代码块"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return para

# 解析Markdown内容
lines = markdown_content.split('\n')
current_section = ""
code_block = False
code_content = ""

for line in lines:
    line = line.rstrip()
    
    if line.startswith('# '):
        # 一级标题
        if code_block:
            add_code_block(doc, code_content)
            code_block = False
            code_content = ""
        add_heading(doc, line[2:], level=0)
    elif line.startswith('## '):
        # 二级标题
        if code_block:
            add_code_block(doc, code_content)
            code_block = False
            code_content = ""
        add_heading(doc, line[3:], level=1)
    elif line.startswith('### '):
        # 三级标题
        if code_block:
            add_code_block(doc, code_content)
            code_block = False
            code_content = ""
        add_heading(doc, line[4:], level=2)
    elif line.startswith('```'):
        # 代码块开始或结束
        if not code_block:
            code_block = True
        else:
            add_code_block(doc, code_content)
            code_block = False
            code_content = ""
    elif code_block:
        # 代码块内容
        code_content += line + '\n'
    elif line.startswith('- '):
        # 列表项
        if code_block:
            add_code_block(doc, code_content)
            code_block = False
            code_content = ""
        para = doc.add_paragraph()
        para.add_run('• ').bold = True
        para.add_run(line[2:])
    else:
        # 普通文本
        if code_block:
            code_content += line + '\n'
        elif line.strip():
            add_paragraph(doc, line)

# 保存Word文档
doc.save('数字化转型指数查询App开发报告.docx')
print("Word文档已生成：数字化转型指数查询App开发报告.docx")
